from docx import Document
import json

def replace_text_preserve_formatting(doc, old_text, new_text):
    """
    Replace `old_text` with `new_text` in the document while preserving all formatting.
    Handles cases where `old_text` spans multiple runs.
    """
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            # Initialize variables to track runs and their text
            runs = paragraph.runs
            text = ''.join(run.text for run in runs)
            
            if old_text in text:
                # Find the start and end positions of the old text
                start = text.find(old_text)
                end = start + len(old_text)
                
                # Clear the paragraph and rebuild it
                paragraph.clear()
                
                # Add text before the old text
                if start > 0:
                    paragraph.add_run(text[:start])
                
                # Add the new text with formatting from the first run of the old text
                new_run = paragraph.add_run(new_text)
                # Find the first run that contains part of the old text
                for run in runs:
                    if run.text in old_text:
                        # Copy formatting from this run
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                        new_run.font.color.rgb = run.font.color.rgb
                        break
                
                # Add text after the old text
                if end < len(text):
                    paragraph.add_run(text[end:])

def replace_text_in_document(resume_path, output_path, edits):
    """
    Replace text in a Word document and save the modified document.
    """
    # Load the document
    doc = Document(resume_path)
    
    # Apply all edits
    for edit_key, edit_value in edits.items():
        old_text = edit_value["to_be_edited"]
        new_text = edit_value["edited"]
        replace_text_preserve_formatting(doc, old_text, new_text)
    
    # Save the modified document
    doc.save(output_path)
